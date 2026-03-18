"""
Document Processing MCP Tools.
Ported from OPN-Agent GenerationAgents.
Handles:
1. Excel file generation (grouped by NPI/Insurance)
2. Word template filling and PDF conversion
3. Folder merging
"""

import os
import shutil
import uuid
import pythoncom
import pandas as pd
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path
from docx import Document
from docx2pdf import convert

from app.core.logging import get_logger
from app.domains.auth.file_validator import file_validator, FileValidationError

logger = get_logger(__name__)

# --- Configuration Constants (Ported from GroupGenerationAgent) ---
COLUMN_MAPPING = {
    'CPT_Description': 'Description of item(s) and/or service(s)',
    'Claim Number': 'Claim Number',
    'ProvOrgNPI': 'Name of provider, facility, or provider of air ambulance services, and National Provider Identifier (NPI)',
    'Date of item(s) or service(s)': 'Date provided',
    'Service code(s)': 'Service code',
    'Initial Payment': 'Initial payment (if no initial payment amount, write N/A)',
    'Offer': 'Offer for total out-of- network rate (including any cost sharing)'
}

ALLOWED_MERGE_EXTENSIONS = {".xls", ".xlsx", ".doc", ".docx", ".pdf"}


async def generate_grouped_excel_files(excel_path: str, output_folder: str) -> Dict[str, Any]:
    """
    Generate grouped Excel files from a master Excel sheet.
    Groups by ProvOrgNPI and InsurancePlanName.
    """
    stats = {"total": 0, "success": 0, "skipped": 0, "failed": 0, "files": [], "errors": []}
    
    try:
        # Validate paths
        safe_excel_path = file_validator.validate_path(excel_path)
        safe_output_folder = file_validator.validate_output_path(output_folder)
        
        # Read Data
        df = pd.read_excel(safe_excel_path)
        df.columns = df.columns.str.strip()
        os.makedirs(safe_output_folder, exist_ok=True)
        

        # Check required columns
        missing_cols = set(COLUMN_MAPPING.keys()) - set(df.columns)
        if "OpenNegGroup" not in df.columns:
            missing_cols.add("OpenNegGroup")
        if "Provider" not in df.columns:
             missing_cols.add("Provider")
        if "InsurancePlanName" not in df.columns:
             missing_cols.add("InsurancePlanName")

        if missing_cols:
             error_msg = f"Missing columns for Group Generation: {missing_cols}"
             logger.warning(error_msg)
             stats["errors"].append(error_msg)
             return stats

        # Unique combinations
        unique_combinations = df[['ProvOrgNPI', 'Provider', 'InsurancePlanName']].drop_duplicates()
        stats["total"] = len(unique_combinations)
        
        logger.info("Starting Excel generation", count=stats["total"])

        for idx, row in unique_combinations.iterrows():
            try:
                # Filter
                filtered_df = df[
                    (df['ProvOrgNPI'] == row['ProvOrgNPI']) &
                    (df['Provider'] == row['Provider']) &
                    (df['InsurancePlanName'] == row['InsurancePlanName'])
                ]
                
                # Validation
                if filtered_df.empty or pd.isna(filtered_df.iloc[0].get('OpenNegGroup')):
                    stats["skipped"] += 1
                    continue
                
                # Transformation
                output_df = filtered_df[list(COLUMN_MAPPING.keys())].rename(columns=COLUMN_MAPPING)
                
                # Format Dates
                output_df['Date provided'] = pd.to_datetime(
                    output_df['Date provided'], errors='coerce'
                ).dt.strftime('%b %d, %Y')
                
                # Add SNO
                output_df.insert(0, 'SNO', range(1, len(output_df) + 1))
                
                # Format Currency
                for col in [
                    'Initial payment (if no initial payment amount, write N/A)',
                    'Offer for total out-of- network rate (including any cost sharing)'
                ]:
                    output_df[col] = output_df[col].apply(_format_currency)

                # Output Path Calculation
                npi_str = _safe_filename(str(row['ProvOrgNPI']))
                insurance_str = str(row['InsurancePlanName']).strip()
                target_dir = safe_output_folder / npi_str / insurance_str
                target_dir.mkdir(parents=True, exist_ok=True)
                
                group_filename = str(filtered_df.iloc[0]['OpenNegGroup']).strip()
                if not group_filename.lower().endswith('.xlsx'):
                    group_filename += '.xlsx'
                    
                full_path = target_dir / group_filename
                
                # Write
                output_df.to_excel(full_path, index=False)
                
                stats["success"] += 1
                stats["files"].append(str(full_path))
                
            except Exception as e:
                stats["failed"] += 1
                error_msg = f"Failed for NPI {row.get('ProvOrgNPI')}: {str(e)}"
                stats["errors"].append(error_msg)
                logger.error(error_msg)
                
        return stats

    except Exception as e:
        logger.error("Global error in generate_grouped_excel_files", error=str(e))
        raise

async def generate_notice_with_pdf(
    excel_path: str, 
    template_path: str, 
    output_folder: str
) -> Dict[str, Any]:
    """
    Generate Word notices and convert to PDF using a template.
    """
    stats = {"total": 0, "success": 0, "skipped": 0, "failed": 0, "files": [], "errors": []}
    
    # Initialize COM for docx2pdf (needed in async context)
    pythoncom.CoInitialize()
    
    try:
        safe_excel_path = file_validator.validate_path(excel_path)
        safe_template_path = file_validator.validate_path(template_path)
        safe_output_folder = file_validator.validate_output_path(output_folder)
        
        # Read Data
        df = pd.read_excel(safe_excel_path, engine='openpyxl')
        df.columns = df.columns.str.strip()
        

        # Check required columns
        required_cols = {"ProvOrgNPI", "Hospital Name", "OpenNegNotice", "InsurancePlanName"}
        missing_cols = required_cols - set(df.columns)
        
        if missing_cols:
             error_msg = f"Missing columns for Notice Generation: {missing_cols}"
             logger.warning(error_msg)
             stats["errors"].append(error_msg)
             return stats

        # Unique combinations
        df_unique = df.drop_duplicates(
            subset=['ProvOrgNPI', 'Hospital Name', 'OpenNegNotice', 'InsurancePlanName']
        )
        stats["total"] = len(df_unique)
        
        logger.info("Starting Notice generation", count=stats["total"])
        
        os.makedirs(safe_output_folder, exist_ok=True)

        for idx, row in df_unique.iterrows():
            try:
                if pd.isna(row.get('OpenNegNotice')):
                    stats["skipped"] += 1
                    continue
                    
                # Load Template
                doc = Document(safe_template_path)
                
                # Replacements
                replacements = {
                    '{Hospital Name}': str(row.get('Hospital Name', '')),
                    '{Provider}': str(row.get('Provider', '')),
                    '{InsurancePlanName}': str(row.get('InsurancePlanName', '')),
                    '{Notice Date}': str(row.get('Notice Date', '')),
                    '{CMS Date1}': str(row.get('CMS Date1', '')),
                    '{CMS Date2}': str(row.get('CMS Date2', '')),
                }
                bold_keys = set(replacements.keys())
                
                # Replace in Paragraphs & Tables
                for paragraph in doc.paragraphs:
                    _replace_placeholders(paragraph, replacements, bold_keys)
                    
                for table in doc.tables:
                    for table_row in table.rows:
                        for cell in table_row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_placeholders(paragraph, replacements, bold_keys)
                
                # Output Paths
                npi_str = str(row['ProvOrgNPI'])
                insurance_str = str(row['InsurancePlanName'])
                target_dir = safe_output_folder / npi_str / insurance_str
                target_dir.mkdir(parents=True, exist_ok=True)
                
                notice_filename = str(row['OpenNegNotice']).strip()
                base_filename = os.path.splitext(notice_filename)[0]
                
                docx_path = target_dir / (base_filename + ".docx")
                pdf_path = target_dir / (base_filename + ".pdf")
                
                # Save Docx
                doc.save(docx_path)
                
                # Convert to PDF
                try:
                    convert(str(docx_path), str(pdf_path))
                    stats["files"].append(str(pdf_path))
                except Exception as e:
                    logger.warning("PDF conversion failed", error=str(e))
                    stats["files"].append(str(docx_path)) # Fallback
                    stats["errors"].append(f"PDF failed for {base_filename}: {e}")
                
                stats["success"] += 1
                
            except Exception as e:
                stats["failed"] += 1
                error_msg = f"Failed for Notice {row.get('OpenNegNotice')}: {str(e)}"
                stats["errors"].append(error_msg)
                logger.error(error_msg)
                
        return stats
        
    except Exception as e:
        logger.error("Global error in generate_notice_with_pdf", error=str(e))
        raise
    finally:
        pythoncom.CoUninitialize()

async def merge_folders(folder1: str, folder2: str, output_folder: str) -> Dict[str, Any]:
    """
    Merge functionality from MergeAgent.
    Consolidates files from two source folders into destination.
    """
    stats = {"copied": 0, "errors": []}
    
    try:
        safe_folder1 = file_validator.validate_path(folder1)
        safe_folder2 = file_validator.validate_path(folder2)
        safe_output = file_validator.validate_output_path(output_folder)
        
        os.makedirs(safe_output, exist_ok=True)
        
        for root_folder in [safe_folder1, safe_folder2]:
            if not root_folder.exists():
                continue
                
            for root, _, files in os.walk(root_folder):
                rel_path = os.path.relpath(root, root_folder)
                dest_subfolder = safe_output / rel_path
                dest_subfolder.mkdir(parents=True, exist_ok=True)
                
                for file in files:
                    ext = os.path.splitext(file)[1].lower()
                    if ext in ALLOWED_MERGE_EXTENSIONS:
                        try:
                            src_path = Path(root) / file
                            _safe_copy(src_path, dest_subfolder)
                            stats["copied"] += 1
                        except Exception as e:
                            stats["errors"].append(f"Failed to copy {file}: {e}")

        logger.info("Merge complete", stats=stats)
        return stats
        
    except Exception as e:
        logger.error("Global error in merge_folders", error=str(e))
        raise

# --- Validation Constants ---
REQUIRED_COLUMNS_GROUP = {
    "ProvOrgNPI", "Provider", "InsurancePlanName", "OpenNegGroup",
    "CPT_Description", "Claim Number", "Date of item(s) or service(s)",
    "Service code(s)", "Initial Payment", "Offer"
}

REQUIRED_COLUMNS_NOTICE = {
    "ProvOrgNPI", "Hospital Name", "Provider", "InsurancePlanName",
    "OpenNegNotice", "Notice Date", "CMS Date1", "CMS Date2"
}

TEMPLATE_PLACEHOLDERS = {
    "{Hospital Name}", "{Provider}", "{InsurancePlanName}",
    "{Notice Date}", "{CMS Date1}", "{CMS Date2}"
}


async def validate_document_request(
    excel_path: str,
    template_path: Optional[str] = None
) -> Dict[str, Any]:
    """
    Validate input data (Excel structure, Template placeholders).
    Returns {"is_valid": bool, "errors": [], "warnings": []}.
    """
    result = {"is_valid": True, "errors": [], "warnings": []}
    
    try:
        # 1. Validate Excel
        safe_excel = file_validator.validate_path(excel_path)
        if not safe_excel.exists():
             result["errors"].append(f"Excel not found: {excel_path}")
             result["is_valid"] = False
             return result

        df = pd.read_excel(safe_excel)
        df.columns = df.columns.str.strip()
        actual_cols = set(df.columns)
        
        # Check combined requirements (Group + Notice if possible)
        # For simplicity, we check if AT LEAST one set is satisfied or partial?
        # Legacy checked based on DocType. Here we check superset?
        # Or just check critical columns.
        
        missing_group = REQUIRED_COLUMNS_GROUP - actual_cols
        missing_notice = REQUIRED_COLUMNS_NOTICE - actual_cols
        
        if missing_group and missing_notice:
            # Missing columns for BOTH types
            result["errors"].append(f"Missing columns for Group: {missing_group}")
            result["errors"].append(f"Missing columns for Notice: {missing_notice}")
            result["is_valid"] = False
            
        # Data Quality
        if "ProvOrgNPI" in df.columns and df["ProvOrgNPI"].isna().sum() > 0:
            result["warnings"].append("Rows with missing ProvOrgNPI found")

        # 2. Validate Template (if provided)
        if template_path:
            safe_template = file_validator.validate_path(template_path)
            if not safe_template.exists():
                result["errors"].append(f"Template not found: {template_path}")
                result["is_valid"] = False
            else:
                doc = Document(safe_template)
                all_text = " ".join([p.text for p in doc.paragraphs])
                for t in doc.tables:
                    for r in t.rows:
                        for c in r.cells:
                            all_text += c.text
                            
                for p in TEMPLATE_PLACEHOLDERS:
                    if p not in all_text:
                        result["warnings"].append(f"Missing placeholder: {p}")

        return result
        
    except Exception as e:
        logger.error("Validation failed", error=str(e))
        result["is_valid"] = False
        result["errors"].append(str(e))
        return result


async def run_document_workflow(
    excel_path: Optional[str] = None,
    template_path: Optional[str] = None,
    output_group_folder: Optional[str] = None,
    output_notice_folder: Optional[str] = None,
    merged_output_folder: Optional[str] = None
) -> Dict[str, Any]:
    """
    Orchestrate the full Phase 2 pipeline:
    Validate -> Generate Groups -> Generate Notices -> Merge.
    """
    overall_stats = {
        "status": "PENDING", 
        "validation": {}, 
        "group_stats": {}, 
        "notice_stats": {}, 
        "merge_stats": {},
        "errors": []
    }
    print(f"DEBUG: run_document_workflow START. Excel={excel_path}, Template={template_path}")
    print(f"DEBUG: Output Group={output_group_folder}")
    
    try:
        # 1. Validate
        val_result = await validate_document_request(excel_path, template_path)
        print(f"DEBUG: Validation Result: {val_result}")
        overall_stats["validation"] = val_result
        if not val_result["is_valid"]:
            print("DEBUG: Validation FAILED.")
            overall_stats["status"] = "FAILED"
            overall_stats["errors"].extend(val_result["errors"])
            return overall_stats

        # 2. Generate Groups
        try:
            print("DEBUG: Calling generate_grouped_excel_files...")
            group_stats = await generate_grouped_excel_files(excel_path, output_group_folder)
            print(f"DEBUG: Group Stats: {group_stats}")
            overall_stats["group_stats"] = group_stats
        except Exception as e:
             print(f"DEBUG: Group Gen Exception: {e}")
             overall_stats["errors"].append(f"Group Gen Failed: {e}")

        # 3. Generate Notices
        if template_path:
            try:
                print("DEBUG: Calling generate_notice_with_pdf...")
                notice_stats = await generate_notice_with_pdf(excel_path, template_path, output_notice_folder)
                print(f"DEBUG: Notice Stats: {notice_stats}")
                overall_stats["notice_stats"] = notice_stats
            except Exception as e:
                 print(f"DEBUG: Notice Gen Exception: {e}")
                 overall_stats["errors"].append(f"Notice Gen Failed: {e}")
        else:
             print("DEBUG: No template path provided, skipping Notice generation.")
             notice_stats = {"success": 0}
             
        # 4. Merge (only if we have some outputs)
        try:
            merge_stats = await merge_folders(output_group_folder, output_notice_folder, merged_output_folder)
            overall_stats["merge_stats"] = merge_stats
        except Exception as e:
             overall_stats["errors"].append(f"Merge Failed: {e}")

        # Final Status
        if overall_stats["errors"]:
            overall_stats["status"] = "PARTIAL_FAILURE" if (group_stats.get("success",0) > 0 or notice_stats.get("success",0) > 0) else "FAILED"
        else:
            overall_stats["status"] = "SUCCESS"
            
        return overall_stats
        
        return overall_stats
        
    except Exception as e:
        print(f"DEBUG: WORKFLOW CRITICAL ERROR: {e}")
        logger.error("Workflow failed", error=str(e))
        overall_stats["status"] = "FAILED"
        overall_stats["errors"].append(str(e))
        return overall_stats

# --- Helper Functions ---

def _format_currency(x) -> str:
    """Format value as currency or N/A."""
    if pd.isna(x) or str(x).strip().upper() == 'N/A':
        return 'N/A'
    try:
        value = float(str(x).replace('$', '').replace(',', '').strip())
        return f"${value:,.2f}"
    except (ValueError, TypeError):
        return str(x)

def _safe_filename(value: str) -> str:
    """Convert value to safe filename."""
    return "".join(c if c.isalnum() else "_" for c in value)

def _safe_copy(src_path: Path, dest_folder: Path) -> Path:
    """Copy file with renaming if exists."""
    base = src_path.stem
    ext = src_path.suffix
    file_name = f"{base}{ext}"
    dest_path = dest_folder / file_name
    
    counter = 1
    while dest_path.exists():
        dest_path = dest_folder / f"{base}_{counter}{ext}"
        counter += 1
    
    shutil.copy2(src_path, dest_path)
    return dest_path

def _replace_placeholders(paragraph, replacements, bold_keys):
    """Replace placeholders in docx paragraph, preserving bold."""
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Quick check
    if not any(k in full_text for k in replacements):
        return
        
    replaced = False
    for k, v in replacements.items():
        if k in full_text:
            full_text = full_text.replace(k, v)
            replaced = True
            
    if replaced:
        # Clear
        for run in paragraph.runs:
            run.text = ""
        # Rebuild
        i = 0
        while i < len(full_text):
            match_found = False
            for k, v in replacements.items():
                if full_text[i:].startswith(v):
                    run = paragraph.add_run(v)
                    if k in bold_keys:
                        run.bold = True
                    i += len(v)
                    match_found = True
                    break
            if not match_found:
                paragraph.add_run(full_text[i])
                i += 1
